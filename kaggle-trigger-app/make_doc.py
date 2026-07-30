import subprocess
import sys
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_documentation():
    doc = docx.Document()

    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("AUTOMATED TRUE CRIME VIDEO GENERATOR\nFrontend Guide & System Credentials")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Dark navy

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("System Architecture, Tokens, Endpoints, and User Manual")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) # Slate Blue
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        return p

    def add_key_value(key, value, is_code=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        run_k = p.add_run(f"• {key}: ")
        run_k.font.bold = True
        run_v = p.add_run(value)
        if is_code:
            run_v.font.name = 'Consolas'
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = RGBColor(0x80, 0x5A, 0xD5) # Purple

    # 1. OVERVIEW
    add_heading_1("1. System Overview")
    p = doc.add_paragraph("The Automated True Crime Video Generator is an end-to-end AI multimedia production pipeline. It connects a modern, 3-tab glassmorphic web frontend (hosted on Vercel) with an automated backend API (hosted on Hugging Face Spaces) that dynamically orchestrates background GPU video rendering inside Kaggle Notebooks (Nvidia Tesla T4).")
    
    p2 = doc.add_paragraph("The 5-Phase Automated Video Pipeline:")
    p2.paragraph_format.space_before = Pt(6)
    p2.paragraph_format.keep_with_next = True
    
    phases = [
        ("Phase 1: Image Generation", "Renders cinematic visuals from script prompts using either local GPU FLUX.1-Schnell or Z Image API (kie.ai)."),
        ("Phase 2: Voice Synthesis", "Generates studio-quality true crime voiceover narration using Kokoro TTS (82M model)."),
        ("Phase 3: Visual & Captions", "Applies cinematic pan/zoom motion and burns word-by-word highlighted captions via FFmpeg."),
        ("Phase 4: Audio-Visual Stitching", "Synchronizes voice narration with matching visual scenes into individual clips."),
        ("Phase 5: Master Concatenation", "Merges all clips, adjusts overall pacing speed, overlays background music (BGM), and exports final MP4.")
    ]
    for ph, desc in phases:
        p_ph = doc.add_paragraph()
        p_ph.paragraph_format.left_indent = Inches(0.3)
        p_ph.paragraph_format.space_after = Pt(3)
        rk = p_ph.add_run(f"✔ {ph}: ")
        rk.font.bold = True
        p_ph.add_run(desc)

    # 2. CREDENTIALS & TOKENS
    add_heading_1("2. System Tokens & API Credentials")
    doc.add_paragraph("Below are the active authentication tokens and API keys required for system operation and deployment:")
    
    add_heading_2("Kaggle API Credentials (GPU Execution Engine)")
    add_key_value("Kaggle Username", "ikechukwuebiringa1", is_code=True)
    add_key_value("Kaggle Access Token (KGAT)", "KGAT_0f12d3a4d07d48f7775e36f82bbc41b6", is_code=True)
    
    add_heading_2("Hugging Face Credentials (Backend API Hosting)")
    add_key_value("Hugging Face Token", "hf_REDACTED", is_code=True)
    add_key_value("Authenticated HF Account", "Airpyk98", is_code=True)

    add_heading_2("Vercel Credentials (Frontend Web Hosting)")
    add_key_value("Vercel CLI Access Token", "vcp_REDACTED", is_code=True)
    add_key_value("Vercel Project ID", "prj_97pMX5nPwnIunnb3gsVwi8mMYkf0", is_code=True)
    add_key_value("Vercel Team Scope", "team_9wXyzLAAUgdkunyH38rnZ20d", is_code=True)

    add_heading_2("Z Image API Engine (kie.ai)")
    add_key_value("API Key Reference", "d392ff55...9f35 (Securely configured inside the Frontend Credentials Tab / localStorage)", is_code=True)

    # 3. ENDPOINTS & REPOSITORIES
    add_heading_1("3. Endpoints & Project Repositories")
    add_key_value("Vercel Production Frontend", "https://kaggle-trigger-app.vercel.app")
    add_key_value("Vercel Aliased Frontend", "https://public-wine-three-41.vercel.app")
    add_key_value("Hugging Face Space Backend", "https://huggingface.co/spaces/epic98/truecrime-video-generator")
    add_key_value("Direct HF API Endpoint", "https://epic98-truecrime-video-generator.hf.space")
    add_key_value("GitHub Source Repository", "https://github.com/Airpyk-98/truecrimescripts.git")

    # 4. FRONTEND USER MANUAL
    add_heading_1("4. Frontend Web App User Guide")
    doc.add_paragraph("The frontend is organized into 3 intuitive glassmorphic tabs. State and logs persist automatically across browser refreshes.")

    add_heading_2("Tab 1: Credentials & Setup")
    doc.add_paragraph("Before initiating runs, verify your API connections in this tab. All values entered here are encrypted and stored locally in your browser's localStorage:")
    add_key_value("Kaggle Authentication", "Input your Kaggle Username and Access Token (KGAT_...). Click 'Test Connection' to verify API reachability.")
    add_key_value("Hugging Face Override", "Optional field to override default HF token for model downloads.")
    add_key_value("Z Image API Key", "Input your kie.ai Bearer token to enable cloud-based image rendering.")

    add_heading_2("Tab 2: Execution & Pipeline Control")
    doc.add_paragraph("This is the main production control center:")
    add_key_value("1. Upload CSV Script", "Select the CSV script file generated by your True Crime AI research skill.")
    add_key_value("2. Aspect Ratio", "Choose between 16:9 (YouTube Landscape), 9:16 (Shorts/Reels/TikTok), or 1:1 (Square).")
    add_key_value("3. Kokoro TTS Voice", "Select narration style. Recommended: 'am_michael' (deep documentary voice) or 'af_heart' (clear storytelling).")
    add_key_value("4. Captions Styling", "Toggle captions on/off. Customize font size (default: 22px), font color (default: white), and black outline thickness.")
    add_key_value("5. Image Engine Selection", "Toggle 'Use Z Image API'. If checked, renders images via cloud API (fast, high consistency). If unchecked, uses local Kaggle GPU FLUX.1-Schnell.")
    add_key_value("6. Trigger Generation", "Click 'Generate Video on Kaggle'. A realtime progress bar and auto-scrolling terminal log will track the 5 phases.")
    
    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent = Inches(0.2)
    rn = p_note.add_run("★ Realtime Monitor Resilience: ")
    rn.font.bold = True
    rn.font.color.rgb = RGBColor(0xD6, 0x9E, 0x2E) # Gold
    p_note.add_run("If you accidentally refresh or close the tab during a 15-minute render, simply reopen the page. The frontend will instantly reconnect to the active job and resume live logs!")

    add_heading_2("Tab 3: Logs & Video Downloads")
    doc.add_paragraph("Review historical productions and export files:")
    add_key_value("Job History Grid", "Displays all past rendering tasks with exact start timestamps and completion durations.")
    add_key_value("Direct MP4 Download", "When a pipeline finishes, a 'Download Video' button appears alongside an embedded video preview player.")
    add_key_value("Log Archiving", "Click 'View Logs' on any past job to inspect the full FFmpeg and Kokoro execution traces.")

    # 5. TROUBLESHOOTING
    add_heading_1("5. Troubleshooting & Error Recovery")
    add_key_value("Cancelled Kaggle Runs", "If you manually cancel a notebook on Kaggle.com, the backend detects the cancel state or 404. The frontend spinner will immediately stop and display 'Pipeline Error / Cancelled'.")
    add_key_value("Z Image 403 Forbidden", "Resolved in v1.2. Image CDN requests are explicitly authorized with Bearer headers to prevent firewall blocks.")
    add_key_value("Disk Space / Temporary Files", "The system automatically purges Kaggle temporary directories 10 seconds after job completion.")

    # Save to Downloads folder
    user_home = os.environ.get('USERPROFILE', r'C:\Users\DELL')
    downloads_dir = os.path.join(user_home, 'Downloads')
    os.makedirs(downloads_dir, exist_ok=True)
    
    output_docx = os.path.join(downloads_dir, "TrueCrime_Pipeline_Documentation.docx")
    doc.save(output_docx)
    print(f"SUCCESS: Documentation generated and saved to: {output_docx}")

if __name__ == "__main__":
    create_documentation()
